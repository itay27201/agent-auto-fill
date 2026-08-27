"""POST /sessions/{session_id}/render

Produces the filled document and returns a presigned download URL.

Three paths:
  pdf_acroform  set real form-field values with pypdf. Highest fidelity —
                the document keeps its structure and stays machine-readable.
  pdf_flat      stamp a text layer at the field bboxes and merge it in.
  docx          replace placeholder runs in the original Word file.

`flatten` matters: some agencies require a flattened PDF that can no longer
be edited, others require the fillable form intact. Offer both.
"""
import io
import logging
import re
import unicodedata
from urllib.parse import quote

from common import config, schema as sch
from common.api import ApiError, body_of, caller, handler, path_param
from common.aws import s3
from common.store import get_session, get_values, load_schema

log = logging.getLogger()
log.setLevel(logging.INFO)


@handler
def lambda_handler(event, _context):
    sid = path_param(event, "session_id")
    body = body_of(event)
    sess = get_session(sid)
    if not sess:
        raise ApiError("session not found", 404)
    if sess.get("owner") not in (caller(event), "anonymous"):
        raise ApiError("forbidden", 403)
    if sess.get("status") != "ready":
        raise ApiError(f"session is {sess.get('status')}, not ready", 409)

    fields = sch.schema_from_list(load_schema(sess["schema_key"]))
    values = get_values(sid)

    result = sch.validate_all(fields, values)
    if body.get("strict", True):
        if not result["ok"]:
            raise ApiError("form has validation errors", 422, **result)
        if result["awaiting_confirmation"]:
            raise ApiError(
                "some values were drafted by the assistant and not confirmed",
                422,
                # The whole result, not just the one list. Every 422 out of this
                # handler is drawn by the same panel, and a body missing the
                # counts leaves it printing "undefined" beside a real blocker.
                # Cheap: it is already computed.
                **result,
            )

    # An uploaded document sits in DocsBucket; a catalog form's master sits in
    # ArtifactsBucket, because DocsBucket expires its entire contents after
    # seven days and a catalog entry has to outlive that.
    src_bucket = sess.get("doc_bucket") or config.DOCS_BUCKET
    src = s3().get_object(Bucket=src_bucket, Key=sess["doc_key"])["Body"].read()
    doc_type = sess.get("doc_type")
    flatten = bool(body.get("flatten", False))

    unplaced: list[str] = []
    if doc_type == "pdf_acroform":
        out, ext = _fill_acroform(src, fields, values, flatten), "pdf"
    elif doc_type == "pdf_flat":
        (out, unplaced), ext = _stamp_overlay(src, fields, values), "pdf"
    elif doc_type == "docx":
        out, ext = _fill_docx(src, fields, values), "docx"
    else:
        raise ApiError(f"cannot render doc_type {doc_type!r}", 500)

    # A strict render promises the exported document matches the form. A value
    # with nowhere to go breaks that promise silently, which is worse than
    # refusing: the person would file a form they believe is complete.
    if unplaced and body.get("strict", True):
        raise ApiError(
            "some filled fields have no known box on the page and would be "
            "dropped from the export — place them in the viewer first",
            422,
            **result,
            unplaced=_unplaced_rows(unplaced, fields),
        )

    key = f"outputs/{sid}/filled.{ext}"
    s3().put_object(
        Bucket=config.ARTIFACTS_BUCKET, Key=key, Body=out,
        ContentType="application/pdf" if ext == "pdf" else
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ServerSideEncryption="aws:kms",
    )
    url = s3().generate_presigned_url(
        "get_object",
        Params={"Bucket": config.ARTIFACTS_BUCKET, "Key": key,
                "ResponseContentDisposition": _disposition(sess.get("filename"), ext)},
        ExpiresIn=config.DOWNLOAD_URL_TTL,
    )
    return {"download_url": url, "key": key, "flattened": flatten,
            "summary": result, "unplaced": _unplaced_rows(unplaced, fields)}


def _unplaced_rows(field_ids: list[str], fields) -> list[dict]:
    """Ids the renderer could not place, carrying their labels.

    Both the 422 and the 200 go through here so the two answers to "what was
    left out" have one shape — the strict refusal and the draft download show
    the same list, and a field id is not something a person filling a form has
    ever seen. Also retires the per-id linear scan the 422 was doing.
    """
    labels = {f.field_id: f.label for f in fields}
    return [{"field_id": fid, "label": labels.get(fid, fid)} for fid in field_ids]


# ------------------------------------------------------------------ filename
# The name the finished document lands under. It is signed into the presigned
# URL rather than sent as a header, so it is built here, before signing — there
# is no adjusting it afterwards without invalidating the signature.

# Anything a filesystem or a header would rather not carry. The name reaching
# here is either a filename the user typed at upload or a catalog entry's
# display name, and neither is checked anywhere on the way.
_UNSAFE = re.compile(r'[\\/:*?"<>|\x00-\x1f]+')
_KNOWN_EXT = re.compile(r"\.(pdf|docx?|png|jpe?g|tiff?)\Z", re.IGNORECASE)


def _stem(raw: str | None) -> str:
    """The document's name, with no extension and nothing that breaks a path.

    Uploads arrive with a real filename; a catalog session carries the form's
    display name, which has no extension at all — which is why the strip is
    conditional on a *known* one rather than "everything after the last dot".
    That version turns `טופס 101 מהדורה 2.5` into `טופס 101 מהדורה 2`.
    """
    name = _UNSAFE.sub(" ", raw or "")
    name = _KNOWN_EXT.sub("", name.strip())
    return re.sub(r"\s+", " ", name).strip(" .")[:80] or "form"


def _ascii_stem(stem: str) -> str:
    """A best-effort ASCII spelling of `stem`, or "" when there is none.

    Latin letters survive NFKD with their accents peeled off. Hebrew has no
    decomposition to ASCII and disappears entirely — which on this product is
    the ordinary case, not the edge one, and is why the caller has a fallback
    name standing by. Restricting the survivors to [A-Za-z0-9._-] is also what
    makes the quoted-string below safe by construction: no quote and no
    backslash can reach it, so there is nothing to escape.
    """
    folded = unicodedata.normalize("NFKD", stem).encode("ascii", "ignore").decode("ascii")
    return re.sub(r"[^A-Za-z0-9._-]+", "-", folded).strip("-._")[:60]


def _disposition(raw_name: str | None, ext: str) -> str:
    """The `Content-Disposition` the finished document arrives with.

    Header values are ASCII and these forms are named in Hebrew, so RFC 6266's
    pair is used: a plain `filename` any client can read, and a `filename*`
    carrying the real UTF-8 name, which every client since IE9 must prefer.

    `attachment` is also the only reason this is a download at all: the anchor's
    `download` attribute is inert on a cross-origin URL, so without this the
    browser navigates away from the session into its own PDF viewer.
    """
    stem = _stem(raw_name)
    fallback = f"{_ascii_stem(stem) or 'form'}-filled.{ext}"
    # safe="" because RFC 5987's ext-value admits only attr-char unescaped, and
    # quote's default safe="/" is not one of them. Everything quote never
    # touches (A-Za-z0-9_.-~) is an attr-char, so this is strictly conforming.
    encoded = quote(f"{stem}-filled.{ext}", safe="")
    return f'attachment; filename="{fallback}"; filename*=UTF-8\'\'{encoded}'


# --------------------------------------------------------------- acroform

def _patch_pypdf_opt():
    """pypdf returns /Opt for choice fields as [[value, label], ...] but
    validates against the raw list. Normalize to values."""
    from pypdf.constants import FieldDictionaryAttributes
    from pypdf.generic import DictionaryObject

    if getattr(DictionaryObject, "_opt_patched", False):
        return
    original = DictionaryObject.get_inherited

    def patched(self, key, default=None):
        result = original(self, key, default)
        if key == FieldDictionaryAttributes.Opt and isinstance(result, list) and all(
            isinstance(v, list) and len(v) == 2 for v in result
        ):
            return [r[0] for r in result]
        return result

    DictionaryObject.get_inherited = patched
    DictionaryObject._opt_patched = True


def _fill_acroform(src: bytes, fields, values, flatten: bool) -> bytes:
    from pypdf import PdfReader, PdfWriter

    _patch_pypdf_opt()
    reader = PdfReader(io.BytesIO(src))
    writer = PdfWriter(clone_from=reader)

    by_page: dict[int, dict] = {}
    for f in fields:
        v = (values.get(f.field_id) or {}).get("value")
        if v in (None, "", []):
            continue
        b = f.backend or {}
        name = b.get("name") or f.field_id

        if f.type == "checkbox":
            # Checkbox "on" states are per-document (/Yes, /1, /On...) —
            # extraction captured the real one; never hardcode "/Yes".
            v = b.get("checked_value", "/Yes") if v else b.get("unchecked_value", "/Off")
        elif f.type == "multiselect" and isinstance(v, list):
            v = ", ".join(str(x) for x in v)
        else:
            v = str(v)

        by_page.setdefault(int(b.get("page", f.page)), {})[name] = v

    for page_no, field_values in by_page.items():
        idx = max(0, page_no - 1)
        if idx < len(writer.pages):
            writer.update_page_form_field_values(
                writer.pages[idx], field_values, auto_regenerate=False
            )

    # Without this, most viewers render the field values as blank.
    writer.set_need_appearances_writer(True)

    if flatten:
        # Drop /AcroForm so the values become static page content.
        try:
            writer._root_object.pop("/AcroForm", None)
        except Exception:
            pass

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


# ---------------------------------------------------------------- overlay

def _stamp_overlay(src: bytes, fields, values) -> tuple[bytes, list[str]]:
    """Draw a transparent text layer per page and merge it onto the original.

    Returns the document and the ids of any filled field it could not place, so
    the caller can report them rather than shipping a document that is quietly
    missing values.
    """
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas

    reader = PdfReader(io.BytesIO(src))
    writer = PdfWriter()
    font = _register_font()

    per_page: dict[int, list] = {}
    skipped: list[str] = []
    for f in fields:
        v = (values.get(f.field_id) or {}).get("value")
        if v in (None, "", []):
            continue
        # No trustworthy box means no stamp. Drawing it anyway would put the
        # value at the top-left corner of page one, on top of whatever the form
        # prints there — which is the failure this whole path exists to stop.
        # It comes back in the response so the caller can say which fields were
        # dropped instead of the document quietly missing them.
        if f.bbox_confidence == "low" or not any(f.bbox):
            skipped.append(f.field_id)
            continue
        per_page.setdefault(f.page, []).append((f, v))

    if skipped:
        log.warning("not stamped, no known box: %s", ", ".join(skipped))

    for i, page in enumerate(reader.pages, start=1):
        items = per_page.get(i)
        if items:
            w = float(page.mediabox.width)
            h = float(page.mediabox.height)
            buf = io.BytesIO()
            c = canvas.Canvas(buf, pagesize=(w, h))
            for f, v in items:
                _draw_field(c, f, v, w, h, font)
            c.save()
            buf.seek(0)
            page.merge_page(PdfReader(buf).pages[0])
        writer.add_page(page)

    out = io.BytesIO()
    writer.write(out)
    return out.getvalue(), skipped


def _draw_field(c, f, value, page_w, page_h, font):
    # A choice printed as a row of tick squares: mark the one that was chosen,
    # not the one the field happens to be anchored on. Without this the value
    # would be written as text into a box a few points across.
    options = (f.backend or {}).get("option_boxes")
    if options:
        _draw_choice(c, f, value, options, page_w, page_h)
        return

    x0, y0, x1, y1 = f.bbox
    # bbox is normalized top-left origin; PDF user space is bottom-left.
    left = x0 * page_w
    right = x1 * page_w
    bottom = (1.0 - y1) * page_h
    top = (1.0 - y0) * page_h
    size = float((f.backend or {}).get("font_size") or min(11.0, max(7.0, (top - bottom) * 0.62)))
    baseline = bottom + (top - bottom - size) / 2 + size * 0.18

    # `backend.mark` comes from the page's own content stream — the geometry pass
    # recognized this box as one of the form's printed tick squares. That outranks
    # `type`, which is the ingest model's opinion and is sometimes "text" on a box
    # a few points across. Writing a string into one of those is illegible ink laid
    # over the form's own printing; an X or a blank are both strictly better.
    if f.type == "checkbox" or (f.backend or {}).get("mark") == "checkbox":
        if _is_ticked(f, value):
            _draw_tick(c, f.bbox, page_w, page_h)
        return

    text = ", ".join(str(x) for x in value) if isinstance(value, list) else str(value)
    rtl = _is_rtl(text)
    c.setFont(font if (rtl and font) else "Helvetica", size)

    # A box printed as separate character cells wants one character per cell. An
    # id number drawn as a single string starts at one edge and drifts out of
    # step with the cells immediately, which on a 9-digit field reads as a
    # different number.
    comb = (f.backend or {}).get("comb")
    if comb and _draw_comb(c, text, comb, page_w, bottom, top, size, rtl):
        return

    if rtl:
        # Hebrew and Arabic need the bidi algorithm applied before drawing;
        # reportlab lays out glyphs strictly left to right, so an unprocessed
        # RTL string comes out reversed. Right-align it too.
        try:
            from bidi.algorithm import get_display
            text = get_display(text)
        except ImportError:
            log.warning("python-bidi is not installed — RTL values will stamp reversed")
        c.drawRightString(right - 2, baseline, text)
    else:
        c.drawString(left + 2, baseline, text)


def _draw_choice(c, f, value, options, page_w, page_h):
    """Tick the square for each chosen option, and nothing otherwise.

    A value that matches none of the printed choices is left unstamped and
    logged. Writing it somewhere would be a guess at which square was meant,
    and a wrongly ticked box on a tax form reads as a deliberate answer.
    """
    chosen = {str(v).strip() for v in (value if isinstance(value, list) else [value])
              if v not in (None, "", True, False)}
    if value is True and len(options) == 1:
        chosen = {str(options[0].get("value"))}

    hit = False
    for option in options:
        if str(option.get("value")).strip() in chosen:
            _draw_tick(c, option["bbox"], page_w, page_h)
            hit = True
    if not hit:
        log.warning("%s: %r is not one of the printed choices — not stamped",
                    f.field_id, value)


# Words that mean "not this one". A checkbox field should hold a boolean, but a
# mistyped one reaches the renderer holding whatever the agent wrote — and on these
# forms that is as often "לא" as "false".
_NEGATIVE = {"לא", "no", "false", "0", "off", "none", "n"}


def _is_ticked(f, value) -> bool:
    """Whether this square gets an X.

    A boolean answers itself. A string is a field that was typed wrong upstream:
    the value still carries the person's answer, so read it rather than discard it,
    but say so in the log — this is the only place a mistyped tick square surfaces.
    """
    if isinstance(value, bool):
        return value
    if value is None:
        return False
    if isinstance(value, list):
        value = ", ".join(str(x) for x in value)

    text = str(value).strip()
    if not text:
        return False

    ticked = _squash(text).lower() not in _NEGATIVE
    log.warning("%s: tick square holds %r, not a boolean — %s. The field is typed "
                "%r and should be checkbox.",
                f.field_id, value, "stamping X" if ticked else "leaving it blank", f.type)
    return ticked


def _squash(text: str) -> str:
    """Letters and digits only — what the form prints around a word is not part of
    the word. Mirrors `ingest_enrich._squash`."""
    return "".join(ch for ch in (text or "") if ch.isalnum())


def _draw_tick(c, bbox, page_w, page_h):
    """An X centred in one of the form's printed squares.

    Sized to the square rather than to the text scale: these are a few points
    across, so the 7pt floor `_draw_field` uses for writing would overflow one.
    Centred on both axes rather than sitting on a text baseline.
    """
    x0, y0, x1, y1 = bbox
    left, right = x0 * page_w, x1 * page_w
    bottom, top = (1.0 - y1) * page_h, (1.0 - y0) * page_h
    box_w, box_h = right - left, top - bottom
    mark = max(3.0, min(11.0, box_h * 0.9, box_w * 0.9))
    c.setFont("Helvetica-Bold", mark)
    c.drawCentredString((left + right) / 2, bottom + (box_h - mark * 0.72) / 2, "X")


def _draw_comb(c, text, comb, page_w, bottom, top, size, rtl) -> bool:
    """One character per printed cell. Returns whether it drew anything.

    Every way this can be unsure returns False and lets the caller draw the value
    as an ordinary string. A value spread across cells that turn out not to line
    up is harder to read than one written straight, and a value silently trimmed
    to the number of cells is a different value.
    """
    xs = comb.get("xs") or []
    cells = len(xs) - 1
    if cells < 1:
        return False

    chars = _comb_chars(text, cells)
    if not chars:
        return False

    # Cells are laid out left to right on the page whichever way the script runs;
    # what changes is which end the value starts from.
    edges = [(xs[i] * page_w, xs[i + 1] * page_w) for i in range(cells)]
    if rtl:
        edges = edges[::-1]

    fit = min(size, (top - bottom) * 0.72,
              min(right - left for left, right in edges) * 0.9)
    c.setFontSize(fit)
    baseline = bottom + (top - bottom - fit) / 2 + fit * 0.18
    for ch, (left, right) in zip(chars, edges):
        c.drawCentredString((left + right) / 2, baseline, ch)
    return True


def _comb_chars(text: str, cells: int) -> list[str] | None:
    """The characters to distribute, or None if they will not fit.

    A date reaches here as `26/02/1996` for eight cells, because the form prints
    the separators between the cells rather than in them. So the punctuation is
    dropped only when keeping it would overflow — never pre-emptively, since a
    field whose cells really do hold a slash should keep it.
    """
    for candidate in ([ch for ch in text if not ch.isspace()],
                      [ch for ch in text if ch.isalnum()]):
        if 0 < len(candidate) <= cells:
            return candidate
    return None


def _is_rtl(text: str) -> bool:
    return any("֐" <= ch <= "ࣿ" for ch in text)


def _register_font():
    """reportlab's built-in fonts have no Hebrew or Arabic glyphs — they
    render as black boxes. Ship a TTF in the layer.

    Failing here is not cosmetic: `_draw_field` falls back to Helvetica, and
    every Hebrew value on the form exports as garbage. It used to fail silently,
    which is how it went unnoticed — the browser viewer draws its own preview
    with system fonts and looks perfectly correct.
    """
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        if "FormAgentRTL" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("FormAgentRTL", config.RTL_FONT_PATH))
        return "FormAgentRTL"
    except Exception:
        log.warning("no RTL font at %s — Hebrew and Arabic values will not render. "
                    "Check the fonts layer is attached to this function.",
                    config.RTL_FONT_PATH)
        return None


# ------------------------------------------------------------------- docx

def _fill_docx(src: bytes, fields, values) -> bytes:
    """Replace {{field_id}} placeholders and matching content-control tags.

    Replacement is done run by run to preserve formatting; python-docx will
    happily drop styling if you rewrite paragraph.text wholesale.
    """
    from docx import Document

    doc = Document(io.BytesIO(src))
    repl = {}
    for f in fields:
        v = (values.get(f.field_id) or {}).get("value")
        if v in (None, [], False):
            v = ""
        elif isinstance(v, list):
            v = ", ".join(str(x) for x in v)
        elif v is True:
            v = "X"
        repl[f"{{{{{f.field_id}}}}}"] = str(v)

    def patch(paragraph):
        for run in paragraph.runs:
            for token, val in repl.items():
                if token in run.text:
                    run.text = run.text.replace(token, val)

    for p in doc.paragraphs:
        patch(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    patch(p)
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                patch(p)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
